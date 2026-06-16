import docx
import sys
import re
import os

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = docx.shared.Pt(11)

with open('Candid_NTCC_Research_Report.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # Check for image
    img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
    if img_match:
        caption = img_match.group(1)
        img_path = img_match.group(2)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=docx.shared.Inches(6.0))
            if caption:
                p = doc.add_paragraph(caption)
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.italic = True
                p.runs[0].font.size = docx.shared.Pt(9)
        continue

    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        p.add_run(line[2:-2]).bold = True
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        # crude handling of inline bold
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                # crude handling of italic
                subparts = re.split(r'(\*.*?\*)', part)
                for subpart in subparts:
                    if subpart.startswith('*') and subpart.endswith('*'):
                        p.add_run(subpart[1:-1]).italic = True
                    else:
                        p.add_run(subpart)

doc.save('Candid_NTCC_Research_Report_v3.docx')
